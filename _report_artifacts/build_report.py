from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(r"E:\OneDrive - 0dpmr\WebPhotocopy")
TEMPLATE = Path(r"E:\OneDrive - 0dpmr\Chuyên đề\Chuyên đề - Hoàng Văn Thiên.docx")
OUT = ROOT / "Chuyên đề - Hoàng Văn Thiên - WebPhotocopyHub - Hoàn thiện.docx"
SHOT_DIR = ROOT / "_report_artifacts" / "screenshots"
DOC_IMG_DIR = ROOT / "_report_artifacts" / "doc_images"

TITLE = "XÂY DỰNG HỆ THỐNG WEB QUẢN LÝ DỊCH VỤ PHOTOCOPY/IN ẤN TRỰC TUYẾN WEBPHOTOCOPYHUB"
STUDENT = "Hoàng Văn Thiên"
STUDENT_UPPER = "HOÀNG VĂN THIÊN"
STUDENT_ID = "225051915"
CLASS_NAME = "22D1ITE-SWE03"
ADVISOR = "Th.S. Hoàng Văn Hiếu"
TABLE_WIDTH_IN = 6.1


def prepare_doc_images() -> dict[str, Path]:
    DOC_IMG_DIR.mkdir(parents=True, exist_ok=True)
    result: dict[str, Path] = {}
    keep_full = {
        "26-architecture-layers.png",
        "27-erd-simplified.png",
        "28-business-flow.png",
        "29-build-result.png",
        "30-use-case-overview.png",
    }
    for src in SHOT_DIR.glob("*.png"):
        if "failed" in src.name:
            continue
        dst = DOC_IMG_DIR / src.name
        img = Image.open(src).convert("RGB")
        if src.name not in keep_full:
            w, h = img.size
            crop_h = min(h, int(w * 0.68), 950)
            img = img.crop((0, 0, w, crop_h))
        img.thumbnail((1500, 1050), Image.Resampling.LANCZOS)
        img.save(dst, "PNG", optimize=True)
        result[src.name] = dst
    return result


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def make_table_header(row):
    for cell in row.cells:
        set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(15, 35, 71)


def mark_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_in: list[float], indent_dxa: int = 0):
    widths_dxa = [str(int(width * 1440)) for width in widths_in]
    total_dxa = str(sum(int(width * 1440) for width in widths_in))
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), total_dxa)

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    old_grid = table._tbl.tblGrid
    if old_grid is not None:
        table._tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), width)
        grid.append(col)
    table._tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), widths_dxa[idx])
            cell.width = Inches(widths_in[idx])


def default_table_widths(col_count: int) -> list[float]:
    defaults = {
        2: [2.35, 3.75],
        3: [1.25, 1.85, 3.0],
        4: [0.55, 1.05, 1.95, 2.55],
        5: [0.85, 1.2, 1.55, 1.55, 0.95],
    }
    return defaults.get(col_count, [TABLE_WIDTH_IN / col_count] * col_count)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\h \\z \\u \\f C \\t "Heading 1,1,Heading 2,2,Heading 3,3"'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_toc_by_tc_field(paragraph, identifier: str):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f"TOC \\h \\z \\f {identifier}"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_tc_field(paragraph, text: str, level: int = 1, identifier: str = "C"):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    escaped = text.replace('"', "'")
    instr.text = f'TC "{escaped}" \\f {identifier} \\l {level}'
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def next_numbering_id(numbering, tag: str) -> int:
    values = []
    for node in numbering.findall(qn(f"w:{tag}")):
        attr = "abstractNumId" if tag == "abstractNum" else "numId"
        value = node.get(qn(f"w:{attr}"))
        if value and value.isdigit():
            values.append(int(value))
    return (max(values) + 1) if values else 1


def ensure_bullet_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering, "abstractNum")
    num_id = next_numbering_id(numbering, "num")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "420")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "420")
    ind.set(qn("w:hanging"), "240")
    p_pr.append(tabs)
    p_pr.append(ind)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Symbol")
    r_fonts.set(qn("w:hAnsi"), "Symbol")
    r_pr.append(r_fonts)
    for node in [start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr]:
        lvl.append(node)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbering_to_paragraph(paragraph, num_id: int, ilvl: int = 0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl_node = num_pr.find(qn("w:ilvl"))
    if ilvl_node is None:
        ilvl_node = OxmlElement("w:ilvl")
        num_pr.append(ilvl_node)
    ilvl_node.set(qn("w:val"), str(ilvl))
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def set_alt_text(paragraph, descr: str):
    for inline in paragraph._p.xpath(".//wp:inline"):
        doc_pr = inline.find(qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", descr)
            doc_pr.set("title", descr[:200])


def ensure_all_image_alt_text(doc: Document):
    for idx, doc_pr in enumerate(doc.element.xpath(".//wp:docPr"), 1):
        if not doc_pr.get("descr"):
            descr = "Hình minh họa trong báo cáo WebPhotocopyHub"
            if idx <= 2:
                descr = "Logo Trường Đại học Kinh tế - Tài chính Thành phố Hồ Chí Minh trên trang bìa"
            elif "Connector" in (doc_pr.get("name") or ""):
                descr = "Đường viền trang bìa theo template báo cáo"
            doc_pr.set("descr", descr)
            doc_pr.set("title", descr[:200])


class ReportBuilder:
    def __init__(self, doc: Document, images: dict[str, Path]):
        self.doc = doc
        self.images = images
        self.table_items: list[str] = []
        self.figure_items: list[str] = []
        self.bullet_num_id = ensure_bullet_numbering(doc)

    def style(self, name: str, fallback: str = "Normal") -> str:
        names = {s.name for s in self.doc.styles}
        return name if name in names else fallback

    def p(self, text: str = "", style: str = "Normal", align=None, bold=False, italic=False):
        paragraph = self.doc.add_paragraph(style=self.style(style))
        if text:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
        if style == "Normal":
            paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
        elif align is not None:
            paragraph.alignment = align
        return paragraph

    def heading(self, text: str, level: int = 1):
        style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
        return self.p(text, style=style)

    def front_heading(self, text: str):
        paragraph = self.p(text, style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
        add_tc_field(paragraph, text, level=1, identifier="C")
        return paragraph

    def front_subheading(self, text: str):
        paragraph = self.p(text, style="Normal", bold=True)
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)
        add_tc_field(paragraph, text, level=2, identifier="C")
        return paragraph

    def title_page_heading(self, text: str):
        p = self.p(text, style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
        return p

    def page_break(self):
        self.doc.add_page_break()

    def bullet(self, text: str):
        p = self.doc.add_paragraph(style=self.style("Normal"))
        add_numbering_to_paragraph(p, self.bullet_num_id)
        p.add_run(text)
        return p

    def table_caption(self, number: str, title: str):
        caption = f"Bảng {number}. {title}"
        self.table_items.append(caption)
        p = self.p(caption, style="Caption", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        add_tc_field(p, caption, level=1, identifier="T")
        return p

    def figure_caption(self, number: str, title: str):
        caption = f"Hình {number}. {title}"
        self.figure_items.append(caption)
        p = self.p(caption, style="Caption", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
        add_tc_field(p, caption, level=1, identifier="F")
        return p

    def add_table(self, number: str, title: str, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
        self.table_caption(number, title)
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.autofit = False
        widths = widths or default_table_widths(len(headers))
        set_table_geometry(table, widths)
        hdr = table.rows[0]
        for i, value in enumerate(headers):
            hdr.cells[i].text = value
        make_table_header(hdr)
        mark_repeat_header(hdr)
        for row in rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = value
                cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cells[i])
                for p in cells[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(value) > 18 else WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(10)
        for cell in hdr.cells:
            set_cell_margins(cell)
        set_table_geometry(table, widths)
        self.p("")
        return table

    def add_figure(self, number: str, title: str, image_name: str, max_width=6.1, max_height=4.5):
        path = self.images[image_name]
        with Image.open(path) as img:
            w, h = img.size
        ratio = w / h
        width = max_width
        height = width / ratio
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if height > max_height:
            run = p.add_run()
            run.add_picture(str(path), height=Inches(max_height))
        else:
            run = p.add_run()
            run.add_picture(str(path), width=Inches(width))
        set_alt_text(p, f"Hình {number}. {title}")
        self.figure_caption(number, title)
        return p


def replace_cover_text(doc: Document):
    title_lines = [
        "XÂY DỰNG HỆ THỐNG WEB QUẢN LÝ DỊCH VỤ",
        "PHOTOCOPY/IN ẤN TRỰC TUYẾN WEBPHOTOCOPYHUB",
    ]
    title_replacements = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "TÊN ĐỀ TÀI":
            clear_paragraph(p)
            r = p.add_run(title_lines[0])
            r.bold = True
            r.add_break()
            r2 = p.add_run(title_lines[1])
            r2.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_replacements += 1
        elif text == "MSSV:225051915":
            clear_paragraph(p)
            p.add_run(f"MSSV: {STUDENT_ID}")
        elif text.startswith("GIẢNG VIÊN HƯỚNG DẪN"):
            clear_paragraph(p)
            p.add_run(f"GIẢNG VIÊN HƯỚNG DẪN: {ADVISOR}")
    if title_replacements < 2:
        raise RuntimeError("Không tìm đủ placeholder TÊN ĐỀ TÀI trên bìa.")


def remove_template_after_second_cover(doc: Document):
    body = doc._body._element
    marker = None
    seen = 0
    for p in doc.paragraphs:
        if "Thành phố Hồ Chí Minh" in p.text:
            seen += 1
            if seen == 2:
                marker = p._p
                break
    if marker is None:
        raise RuntimeError("Không tìm được điểm kết thúc trang bìa phụ.")
    removing = False
    for child in list(body):
        if child is marker:
            removing = True
            continue
        if removing and child.tag != qn("w:sectPr"):
            body.remove(child)


def add_references(builder: ReportBuilder):
    builder.front_heading("TÀI LIỆU THAM KHẢO")
    refs = [
        ("American Psychological Association. (n.d.). Webpage on a website references. APA Style. ",
         "https://apastyle.apa.org/style-grammar-guidelines/references/examples/webpage-website-references"),
        ("Bootstrap team. (n.d.). Get started with Bootstrap. Bootstrap v5.3 documentation. ",
         "https://getbootstrap.com/docs/5.3/getting-started/introduction/"),
        ("Microsoft. (n.d.). Get started with Swashbuckle and ASP.NET Core. Microsoft Learn. ",
         "https://learn.microsoft.com/en-us/aspnet/core/tutorials/getting-started-with-swashbuckle"),
        ("Microsoft. (n.d.). Introduction to authorization in ASP.NET Core. Microsoft Learn. ",
         "https://learn.microsoft.com/en-us/aspnet/core/security/authorization/introduction"),
        ("Microsoft. (n.d.). Introduction to Identity on ASP.NET Core. Microsoft Learn. ",
         "https://learn.microsoft.com/en-us/aspnet/core/security/authentication/identity"),
        ("Microsoft. (n.d.). Overview of ASP.NET Core MVC. Microsoft Learn. ",
         "https://learn.microsoft.com/en-us/aspnet/core/mvc/overview"),
        ("Microsoft. (n.d.). Overview of Entity Framework Core. Microsoft Learn. ",
         "https://learn.microsoft.com/en-us/ef/core/"),
        ("OpenAPI Initiative. (n.d.). OpenAPI Specification. ",
         "https://swagger.io/specification/"),
        ("OWASP Foundation. (n.d.). OWASP Application Security Verification Standard. ",
         "https://owasp.org/www-project-application-security-verification-standard/"),
        ("PostgreSQL Global Development Group. (n.d.). Transaction isolation. PostgreSQL Documentation. ",
         "https://www.postgresql.org/docs/current/transaction-iso.html"),
        ("Purdue Online Writing Lab. (n.d.). In-text citations: The basics. Purdue OWL. ",
         "https://owl.purdue.edu/owl/research_and_citation/apa_style/apa_formatting_and_style_guide/in_text_citations_the_basics.html"),
        ("Supabase. (n.d.). Connect to your database. Supabase Docs. ",
         "https://supabase.com/docs/guides/database/connecting-to-postgres"),
        ("Trường Đại học Kinh tế - Tài chính Thành phố Hồ Chí Minh. (2015). Hướng dẫn viết chuyên đề/khóa luận tốt nghiệp. ",
         "https://www.uef.edu.vn/newsimg/ktckt/UEF_TC-KT_HuongdanvietKLTN_2015.09.12.pdf"),
    ]
    for text, url in refs:
        p = builder.p("", style="Normal", align=WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.add_run(text)
        add_hyperlink(p, url, url)


def build_document():
    images = prepare_doc_images()
    shutil.copyfile(TEMPLATE, OUT)
    doc = Document(OUT)
    replace_cover_text(doc)
    remove_template_after_second_cover(doc)

    builder = ReportBuilder(doc, images)

    # Front matter
    builder.page_break()
    builder.front_heading("LỜI CẢM ƠN")
    for text in [
        f"Em xin gửi lời cảm ơn chân thành đến {ADVISOR}, người đã hướng dẫn, góp ý và định hướng để em hoàn thiện chuyên đề tốt nghiệp này. Những góp ý của thầy giúp em nhìn rõ hơn yêu cầu của một hệ thống web thực tế, đặc biệt là cách tổ chức kiến trúc, bảo mật và kiểm thử sản phẩm.",
        "Em cũng xin cảm ơn quý thầy cô Khoa Công nghệ thông tin, Trường Đại học Kinh tế - Tài chính Thành phố Hồ Chí Minh, đã trang bị nền tảng kiến thức về lập trình web, cơ sở dữ liệu, công nghệ phần mềm và kiểm thử. Các kiến thức này được vận dụng trực tiếp trong quá trình xây dựng hệ thống WebPhotocopyHub.",
        "Cuối cùng, em xin cảm ơn gia đình, bạn bè và những người đã hỗ trợ em trong quá trình thực hiện đề tài. Báo cáo này là kết quả tổng hợp từ quá trình nghiên cứu, xây dựng, kiểm thử và hoàn thiện sản phẩm theo đúng phạm vi chuyên đề tốt nghiệp.",
    ]:
        builder.p(text)
    builder.p(STUDENT, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)

    builder.page_break()
    builder.front_heading("LỜI CAM ĐOAN")
    for text in [
        f"Em tên là {STUDENT}, MSSV {STUDENT_ID}, cam đoan chuyên đề tốt nghiệp với đề tài \"{TITLE.title()}\" là kết quả thực hiện của bản thân dưới sự hướng dẫn của giảng viên hướng dẫn.",
        "Các nội dung phân tích, thiết kế, xây dựng và kiểm thử trong báo cáo được thực hiện dựa trên project WebPhotocopyHub hiện có trong thư mục mã nguồn. Những tài liệu, hình ảnh, lý thuyết và nguồn tham khảo bên ngoài được trích dẫn theo chuẩn APA 7 và liệt kê trong mục Tài liệu tham khảo.",
        "Em chịu trách nhiệm về tính trung thực của nội dung báo cáo, hình ảnh minh họa và kết quả kiểm thử được trình bày.",
    ]:
        builder.p(text)
    builder.p(f"TP. Hồ Chí Minh, năm 2026", align=WD_ALIGN_PARAGRAPH.RIGHT)
    builder.p("Sinh viên thực hiện", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
    builder.p(STUDENT, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)

    builder.page_break()
    builder.front_heading("MỤC LỤC")
    add_toc_field(doc.add_paragraph())

    builder.page_break()
    builder.front_heading("DANH MỤC CÁC KÍ HIỆU, CHỮ VIẾT TẮT")
    builder.add_table(
        "0.1",
        "Danh mục chữ viết tắt sử dụng trong báo cáo",
        ["STT", "Chữ viết tắt", "Từ gốc tiếng Anh", "Diễn giải"],
        [
            ["1", "API", "Application Programming Interface", "Giao diện lập trình ứng dụng"],
            ["2", "APA", "American Psychological Association", "Chuẩn trích dẫn tài liệu tham khảo"],
            ["3", "ASVS", "Application Security Verification Standard", "Chuẩn kiểm chứng bảo mật ứng dụng của OWASP"],
            ["4", "CSDL", "Database", "Cơ sở dữ liệu"],
            ["5", "DTO", "Data Transfer Object", "Đối tượng truyền dữ liệu giữa các lớp"],
            ["6", "EF Core", "Entity Framework Core", "ORM truy cập dữ liệu trong .NET"],
            ["7", "ERD", "Entity Relationship Diagram", "Sơ đồ quan hệ thực thể"],
            ["8", "HTTP/HTTPS", "Hypertext Transfer Protocol/Secure", "Giao thức truyền tải web"],
            ["9", "MVC", "Model-View-Controller", "Mẫu kiến trúc Model-View-Controller"],
            ["10", "QR", "Quick Response", "Mã phản hồi nhanh, định hướng phát triển thanh toán"],
        ],
    )

    builder.page_break()
    builder.front_heading("DANH MỤC CÁC BẢNG")
    add_toc_by_tc_field(doc.add_paragraph(), "T")

    builder.page_break()
    builder.front_heading("DANH MỤC CÁC HÌNH VẼ, ĐỒ THỊ")
    add_toc_by_tc_field(doc.add_paragraph(), "F")

    # Main content
    builder.page_break()
    builder.front_heading("MỞ ĐẦU")
    builder.front_subheading("Lý do chọn đề tài")
    builder.p("Nhu cầu in ấn, photocopy và hoàn thiện tài liệu của sinh viên, nhân viên văn phòng và các đơn vị nhỏ vẫn diễn ra hằng ngày. Tuy nhiên, quy trình truyền thống thường phụ thuộc vào trao đổi trực tiếp, gửi file rời rạc qua nhiều kênh, báo giá thủ công và khó theo dõi trạng thái xử lý. Điều này làm tăng thời gian chờ, dễ nhầm file, khó kiểm soát thanh toán và thiếu dữ liệu để đối soát.")
    builder.p("Đề tài WebPhotocopyHub được lựa chọn nhằm xây dựng một hệ thống web hỗ trợ khách hàng gửi file, tạo đơn in, theo dõi đơn, nạp ví và mua sản phẩm/dịch vụ hỗ trợ; đồng thời cung cấp khu vận hành cho tiệm photocopy và khu quản trị hệ thống cho admin. Đề tài gắn với chuyên ngành Kỹ thuật phần mềm vì yêu cầu phân tích nghiệp vụ, thiết kế kiến trúc, tổ chức cơ sở dữ liệu, bảo mật, kiểm thử và đánh giá sản phẩm thực tế.")
    builder.front_subheading("Mục tiêu đề tài")
    for item in [
        "Xây dựng hệ thống web quản lý dịch vụ photocopy/in ấn theo ba nhóm vai trò: Customer, ShopOperator và Admin.",
        "Chuẩn hóa luồng đặt in: upload file, chọn cấu hình in, tính giá, theo dõi trạng thái và xử lý thanh toán qua ví.",
        "Xây dựng chức năng vận hành cho tiệm: hàng chờ in, duyệt nạp tiền, nạp tại quầy, quản lý tồn kho và đơn dịch vụ.",
        "Xây dựng chức năng quản trị hệ thống: người dùng, bảng giá, sản phẩm, dịch vụ, audit log, đối soát ví, Swagger và health monitoring.",
        "Kiểm thử solution, chụp giao diện thật và đánh giá kết quả theo yêu cầu chuyên đề tốt nghiệp.",
    ]:
        builder.bullet(item)
    builder.front_subheading("Phạm vi và phương pháp thực hiện")
    builder.p("Phạm vi báo cáo tập trung vào phiên bản hiện tại của project WebPhotocopyHub. Các chức năng QR payment và Document Hub được trình bày như hướng phát triển, không mô tả như chức năng đã hoàn thiện. Phương pháp thực hiện gồm khảo sát nhu cầu nghiệp vụ, phân tích yêu cầu, thiết kế kiến trúc và dữ liệu, xây dựng bằng ASP.NET Core MVC, kiểm thử luồng chính và đối chiếu kết quả với project thực tế.")
    builder.front_subheading("Cấu trúc báo cáo")
    builder.p("Báo cáo gồm bảy chương: tổng quan bài toán, cơ sở lý thuyết và công nghệ, phân tích yêu cầu, thiết kế hệ thống, xây dựng chức năng, kiểm thử đánh giá và kết luận hướng phát triển.")

    builder.page_break()
    builder.heading("CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI VÀ BÀI TOÁN", 1)
    builder.heading("Bối cảnh nghiệp vụ", 2)
    builder.p("Các tiệm photocopy phục vụ sinh viên và văn phòng thường xử lý nhiều loại yêu cầu: in tài liệu, photocopy hồ sơ, đóng gáy, scan, mua văn phòng phẩm và giao nhận tài liệu. Khi các yêu cầu được tiếp nhận trực tiếp hoặc qua tin nhắn, nhân viên phải tự ghi chú file, số bản, in màu, khổ giấy và thời gian nhận. Nếu không có hệ thống quản lý tập trung, file có thể bị thất lạc, yêu cầu dễ bị hiểu sai và người quản lý khó nắm tổng quan công việc.")
    builder.add_figure("1.1", "Trang chủ giới thiệu nền tảng WebPhotocopyHub", "01-public-home.png")
    builder.heading("Vấn đề của quy trình thủ công", 2)
    for item in [
        "Khách hàng khó biết trạng thái đơn sau khi gửi file.",
        "Tiệm photocopy khó phân công, ưu tiên và ghi nhận lịch sử xử lý đơn.",
        "Thanh toán/nạp tiền nếu ghi nhận thủ công dễ lệch số dư và khó đối soát.",
        "Admin thiếu công cụ quản trị người dùng, bảng giá, audit log và giám sát hệ thống.",
    ]:
        builder.bullet(item)
    builder.heading("Mục tiêu giải pháp WebPhotocopyHub", 2)
    builder.p("WebPhotocopyHub hướng đến mô hình quản lý tập trung: mỗi cơ sở có trang khách hàng và khu vận hành riêng; khách hàng tạo đơn và theo dõi online; ShopOperator xử lý hàng chờ; Admin quản trị dữ liệu toàn hệ thống. Giải pháp ưu tiên tính rõ ràng trong quy trình, khả năng truy vết và kiểm soát dữ liệu tài chính.")
    builder.add_figure("1.2", "Trang cơ sở Toàn Photocopy dành cho khách hàng", "03-public-branch.png")

    builder.page_break()
    builder.heading("CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG", 1)
    builder.heading("ASP.NET Core MVC và Razor Views", 2)
    builder.p("ASP.NET Core MVC là framework xây dựng ứng dụng web theo mẫu Model-View-Controller, phù hợp với các hệ thống cần tách lớp xử lý request, giao diện và dữ liệu hiển thị (Microsoft, n.d.). Trong project, controllers điều phối request, Razor Views hiển thị giao diện, còn nghiệp vụ chính được đẩy xuống Application/Infrastructure để tránh phụ thuộc giao diện.")
    builder.heading("Identity, Authorization và bảo mật web", 2)
    builder.p("ASP.NET Core Identity hỗ trợ đăng nhập, mật khẩu, hồ sơ người dùng, role, claim và token (Microsoft, n.d.). WebPhotocopyHub dùng Identity kết hợp role-based authorization và policy authorization để tách quyền Customer, ShopOperator và Admin. Authorization là bước kiểm tra người dùng có quyền thực hiện thao tác sau khi đã xác thực danh tính (Microsoft, n.d.).")
    builder.p("Bên cạnh đó, hệ thống cấu hình anti-forgery, security headers, rate limiting cho auth/money và kiểm soát file upload. Các biện pháp này phù hợp với định hướng của OWASP ASVS về việc chuẩn hóa kiểm soát bảo mật ứng dụng web (OWASP Foundation, n.d.).")
    builder.heading("EF Core, PostgreSQL/Supabase và giao dịch dữ liệu", 2)
    builder.p("Entity Framework Core là ORM giúp lập trình viên .NET làm việc với cơ sở dữ liệu bằng đối tượng .NET và giảm lượng code truy cập dữ liệu thủ công (Microsoft, n.d.). Project dùng EF Core với PostgreSQL/Supabase để lưu user, ví, đơn in, sản phẩm, dịch vụ và audit log. Với nghiệp vụ ví, hệ thống dùng transaction, idempotency key và unique index để hạn chế ghi trùng giao dịch.")
    builder.p("PostgreSQL hỗ trợ các mức cô lập giao dịch, trong đó Serializable giúp giảm rủi ro bất nhất ở các thao tác tài chính nhạy cảm (PostgreSQL Global Development Group, n.d.). Khi dùng Supabase trong môi trường phát triển, connection string qua pooler phù hợp với điều kiện mạng IPv4 và cách vận hành hosted PostgreSQL (Supabase, n.d.).")
    builder.heading("Bootstrap, OpenAPI/Swagger và chuẩn trích dẫn", 2)
    builder.p("Bootstrap 5.3 hỗ trợ xây dựng giao diện responsive bằng hệ thống layout và components có sẵn (Bootstrap team, n.d.). Swagger/OpenAPI giúp mô tả API theo chuẩn có thể đọc bởi người và công cụ, hỗ trợ kiểm thử và tài liệu hóa endpoint (OpenAPI Initiative, n.d.). Trong báo cáo, các nguồn tham khảo được trình bày theo APA 7; trích dẫn trong văn bản theo phương pháp author-date (Purdue Online Writing Lab, n.d.).")
    builder.add_table(
        "2.1",
        "Công nghệ sử dụng trong project WebPhotocopyHub",
        ["Nhóm", "Công nghệ", "Vai trò trong hệ thống"],
        [
            ["Backend", "ASP.NET Core MVC .NET 8", "Xử lý request, routing, controller, Razor View và middleware."],
            ["Xác thực", "ASP.NET Core Identity", "Quản lý tài khoản, mật khẩu, role Customer/ShopOperator/Admin."],
            ["Dữ liệu", "EF Core + PostgreSQL/Supabase", "Lưu trữ dữ liệu nghiệp vụ và truy vấn qua DbContext/service."],
            ["Giao diện", "Razor Views, Bootstrap, CSS riêng", "Xây dựng UI public, customer, shop và admin."],
            ["API", "Swagger/OpenAPI", "Tài liệu hóa và kiểm thử nhanh API nội bộ."],
            ["Bảo mật", "Anti-forgery, security headers, rate limiting", "Giảm rủi ro CSRF, cấu hình header và abuse request."],
        ],
    )
    builder.add_figure("2.1", "Kiến trúc 4 lớp của WebPhotocopyHub", "26-architecture-layers.png")

    builder.page_break()
    builder.heading("CHƯƠNG 3. PHÂN TÍCH YÊU CẦU HỆ THỐNG", 1)
    builder.heading("Tác nhân sử dụng", 2)
    builder.add_table(
        "3.1",
        "Nhóm tác nhân và mục tiêu sử dụng",
        ["Tác nhân", "Mục tiêu", "Chức năng chính"],
        [
            ["Customer", "Đặt in và theo dõi yêu cầu tại cơ sở photocopy.", "Đăng nhập, tạo đơn in, nạp ví, mua sản phẩm, đặt dịch vụ, xem lịch sử."],
            ["ShopOperator", "Vận hành công việc tại quầy.", "Xử lý hàng chờ in, duyệt nạp tiền, nạp tại quầy, quản lý tồn kho và đơn."],
            ["Admin", "Quản trị toàn hệ thống.", "Quản lý người dùng, bảng giá, giao dịch ví, audit log, đối soát, giám sát API."],
        ],
    )
    builder.add_figure("3.1", "Use case tổng quát của hệ thống WebPhotocopyHub", "30-use-case-overview.png")
    builder.heading("Yêu cầu chức năng", 2)
    builder.add_table(
        "3.2",
        "Yêu cầu chức năng chính",
        ["Mã", "Nhóm chức năng", "Mô tả"],
        [
            ["FR01", "Tài khoản", "Cho phép đăng ký/đăng nhập, phân quyền theo role và điều hướng theo khu vực."],
            ["FR02", "Đơn in", "Khách hàng upload file, chọn cấu hình in, tạo đơn và theo dõi trạng thái."],
            ["FR03", "Định giá", "Tính giá theo khổ giấy, số trang, số bản, in màu/trắng đen, một mặt/hai mặt."],
            ["FR04", "Ví/nạp tiền", "Tạo yêu cầu nạp, duyệt nạp, ghi WalletTransaction và đối soát số dư."],
            ["FR05", "Vận hành shop", "Nhân viên xử lý đơn in, đơn sản phẩm, dịch vụ hỗ trợ và tồn kho."],
            ["FR06", "Quản trị", "Admin quản lý dữ liệu hệ thống, audit log, báo cáo CSV, health và Swagger."],
        ],
    )
    builder.heading("Yêu cầu phi chức năng", 2)
    for item in [
        "Bảo mật: chỉ người dùng đúng role mới truy cập được khu vực tương ứng.",
        "Toàn vẹn dữ liệu: giao dịch ví và đơn hàng cần idempotency key, unique index và transaction.",
        "Khả dụng: có health endpoint để kiểm tra trạng thái ứng dụng và database.",
        "Dễ bảo trì: tách module Web.Customer, Web.Shop, Web.Admin và các lớp Domain/Application/Infrastructure.",
        "Dễ dùng: giao diện thao tác rõ ràng cho khách hàng và nhân viên vận hành.",
    ]:
        builder.bullet(item)
    builder.add_figure("3.2", "Luồng nghiệp vụ đặt in và thanh toán ví", "28-business-flow.png")

    builder.page_break()
    builder.heading("CHƯƠNG 4. THIẾT KẾ HỆ THỐNG", 1)
    builder.heading("Thiết kế kiến trúc", 2)
    builder.p("Solution WebPhotocopyHub được tổ chức theo hướng phân lớp. Domain chứa entity, enum và constant. Application chứa contract, DTO và chính sách dùng chung. Infrastructure triển khai DbContext, service, cache, routine và report. Web là host ASP.NET Core đồng thời tích hợp module Customer, Shop và Admin. Cách chia này giúp giao diện không chứa trực tiếp nghiệp vụ tài chính và giúp service có thể được kiểm thử/duy trì độc lập hơn.")
    builder.add_table(
        "4.1",
        "Vai trò các module trong solution",
        ["Module", "Vai trò"],
        [
            ["WebPhotocopyHub.Domain", "Định nghĩa entity, enum, role constants và base model."],
            ["WebPhotocopyHub.Application", "Định nghĩa service contracts, DTOs, policy và shared view models."],
            ["WebPhotocopyHub.Infrastructure", "Triển khai EF Core, service nghiệp vụ, storage, routine, cache, audit và report."],
            ["WebPhotocopyHub.Web.Customer", "Giao diện khách hàng theo cơ sở: dashboard, đơn in, ví, sản phẩm, dịch vụ."],
            ["WebPhotocopyHub.Web.Shop", "Khu vận hành cơ sở: dashboard, hàng chờ, nạp tiền tại quầy, tồn kho."],
            ["WebPhotocopyHub.Web.Admin", "Khu quản trị hệ thống: user, bảng giá, audit, đối soát, monitoring."],
            ["WebPhotocopyHub.Web", "Host, routing, middleware, security headers, Swagger và health checks."],
        ],
    )
    builder.heading("Thiết kế dữ liệu", 2)
    builder.p("Cơ sở dữ liệu được thiết kế quanh ApplicationUser và các nhóm nghiệp vụ: ví/nạp tiền, file upload/đơn in, sản phẩm/đơn sản phẩm, dịch vụ hỗ trợ và audit log. Các bảng có index theo trạng thái, người dùng và thời gian để phục vụ truy vấn dashboard, lịch sử giao dịch và danh sách xử lý.")
    builder.add_figure("4.1", "ERD rút gọn các thực thể nghiệp vụ chính", "27-erd-simplified.png")
    builder.add_table(
        "4.2",
        "Một số thực thể dữ liệu trọng tâm",
        ["Thực thể", "Ý nghĩa", "Điểm thiết kế đáng chú ý"],
        [
            ["ApplicationUser", "Người dùng hệ thống.", "Lưu số dư ví, trạng thái hoạt động, thông tin cá nhân."],
            ["WalletTransaction", "Dòng giao dịch ví.", "Có BalanceBefore/After và idempotency key để đối soát."],
            ["TopUpRequest", "Yêu cầu nạp tiền.", "Hỗ trợ duyệt, từ chối và duyệt admin bước 2."],
            ["PrintJob", "Đơn in.", "Lưu cấu hình in, trạng thái, giá tiền, file và thông tin thanh toán."],
            ["UploadedFileMetadata", "Metadata file tải lên.", "Lưu ngoài wwwroot, gắn owner để kiểm soát quyền."],
            ["AuditLog", "Nhật ký hoạt động.", "Có PreviousHash/RecordHash để kiểm tra chuỗi audit."],
        ],
    )
    builder.heading("Thiết kế bảo mật và giao dịch", 2)
    builder.p("Hệ thống dùng global authenticated-user filter, sau đó mở AllowAnonymous cho các trang public/login cần thiết. Với các nghiệp vụ nhạy cảm, policy và role được kiểm tra ở controller. Các thao tác tiền ví dùng transaction, kiểm tra số dư không âm, idempotency key và audit log. Thiết kế này giúp giảm lỗi ghi trùng khi người dùng gửi lại request hoặc khi mạng chập chờn.")

    builder.page_break()
    builder.heading("CHƯƠNG 5. XÂY DỰNG VÀ TRIỂN KHAI CHỨC NĂNG", 1)
    builder.heading("Public site và trang cơ sở", 2)
    builder.p("Nhóm public site giúp người dùng hiểu nhanh dịch vụ, chọn cơ sở và đi vào luồng đăng nhập/đặt in. Trang danh sách cơ sở và landing cơ sở Toàn Photocopy được thiết kế để làm rõ thông tin dịch vụ, giờ mở cửa và hành động chính.")
    builder.add_figure("5.1", "Danh sách cơ sở photocopy trên public site", "02-public-shops.png")
    builder.heading("Customer portal", 2)
    builder.p("Customer portal là khu vực khách hàng sau đăng nhập. Dashboard hiển thị trạng thái đơn, số dư ví và các lối tắt thao tác. Chức năng tạo đơn in cho phép chọn file, nhập cấu hình in và ghi chú. Các màn hình ví, nạp tiền, sản phẩm và dịch vụ hỗ trợ giúp khách hàng thực hiện các nhu cầu thường gặp tại tiệm photocopy.")
    builder.add_figure("5.2", "Dashboard khách hàng", "05-customer-dashboard.png")
    builder.add_figure("5.3", "Màn hình tạo đơn in mới", "06-customer-print-create.png")
    builder.add_figure("5.4", "Màn hình ví và lịch sử giao dịch", "08-customer-wallet.png")
    builder.add_figure("5.5", "Danh sách sản phẩm văn phòng phẩm", "10-customer-products.png")
    builder.heading("Shop operation", 2)
    builder.p("Khu vận hành cơ sở được thiết kế cho nhân viên xử lý công việc trong ngày. Dashboard gom các chỉ số quan trọng như hàng chờ in, yêu cầu nạp tiền, đơn sản phẩm, đơn dịch vụ và cảnh báo tồn kho. Nhân viên có thể vào hàng chờ in, duyệt nạp tiền hoặc nạp trực tiếp tại quầy.")
    builder.add_figure("5.6", "Dashboard vận hành của ShopOperator", "12-shop-dashboard.png")
    builder.add_figure("5.7", "Hàng chờ in của cơ sở", "13-shop-print-queue.png")
    builder.add_figure("5.8", "Màn hình nạp tiền tại quầy", "15-shop-counter-topup.png")
    builder.add_figure("5.9", "Quản lý tồn kho văn phòng phẩm", "16-shop-inventory.png")
    builder.heading("Admin system", 2)
    builder.p("Admin system tập trung vào quản trị toàn hệ thống: người dùng, bảng giá, sản phẩm/dịch vụ, yêu cầu nạp tiền, giao dịch ví, đối soát và giám sát. System monitoring hiển thị trạng thái health check, runtime, memory, API endpoints và audit chain để hỗ trợ vận hành.")
    builder.add_figure("5.10", "Dashboard quản trị hệ thống", "17-admin-dashboard.png")
    builder.add_figure("5.11", "Quản lý người dùng toàn hệ thống", "18-admin-users.png")
    builder.add_figure("5.12", "Trung tâm giám sát hệ thống", "23-admin-monitoring.png")
    builder.heading("API, Swagger và health monitoring", 2)
    builder.p("Swagger UI được bật trong môi trường phát triển để kiểm tra các endpoint như catalog, pricing, me và system. Health endpoints gồm live, ready và database giúp kiểm tra nhanh trạng thái ứng dụng trong quá trình demo hoặc vận hành.")
    builder.add_figure("5.13", "Swagger UI của WebPhotocopyHub API", "24-swagger.png")

    builder.page_break()
    builder.heading("CHƯƠNG 6. KIỂM THỬ VÀ ĐÁNH GIÁ KẾT QUẢ", 1)
    builder.heading("Môi trường và dữ liệu kiểm thử", 2)
    builder.p("Solution được build trên .NET 8, chạy local tại http://localhost:5250 và kết nối PostgreSQL/Supabase qua biến môi trường. Dữ liệu seed gồm tài khoản Admin, ShopOperator, Customer; bảng giá in; sản phẩm; dịch vụ hỗ trợ; đơn in; yêu cầu nạp tiền và giao dịch ví mẫu.")
    builder.add_figure("6.1", "Kết quả build solution WebPhotocopyHub", "29-build-result.png")
    builder.heading("Ca kiểm thử chính", 2)
    builder.add_table(
        "6.1",
        "Bảng kiểm thử các luồng chức năng chính",
        ["STT", "Luồng kiểm thử", "Dữ liệu/điều kiện", "Kết quả mong đợi", "Trạng thái"],
        [
            ["1", "Build solution", "dotnet build .\\WebPhotocopyHub.sln --no-restore", "Tất cả project build thành công, 0 warning, 0 error.", "Đạt"],
            ["2", "Đăng nhập Customer", "Tài khoản khách hàng seed", "Vào được dashboard khách hàng theo cơ sở.", "Đạt"],
            ["3", "Tạo đơn in", "File hợp lệ và cấu hình in hợp lệ", "Tạo PrintJob trạng thái Submitted và tính giá.", "Đạt"],
            ["4", "Xem ví/nạp tiền", "Tài khoản có lịch sử ví", "Hiển thị số dư, giao dịch và form nạp tiền.", "Đạt"],
            ["5", "Đăng nhập ShopOperator", "Tài khoản operator seed", "Vào được khu vận hành cơ sở và xem hàng chờ.", "Đạt"],
            ["6", "Nạp tiền tại quầy", "ShopOperator + khách hàng hợp lệ", "Tạo yêu cầu nạp CounterCash và cập nhật ví.", "Đạt"],
            ["7", "Đăng nhập Admin", "Tài khoản admin seed", "Vào được dashboard admin và các chức năng quản trị.", "Đạt"],
            ["8", "Health DB", "/healthz/db", "Trả trạng thái Healthy khi database sẵn sàng.", "Đạt"],
        ],
    )
    builder.add_figure("6.2", "Health check database trả trạng thái Healthy", "25-health-db.png")
    builder.heading("Đánh giá kết quả", 2)
    builder.p("Kết quả kiểm thử cho thấy project đã hoàn thiện các luồng cốt lõi: khách hàng đặt in và quản lý ví, nhân viên vận hành xử lý đơn và nạp tiền tại quầy, admin quản trị và giám sát hệ thống. Cấu trúc solution rõ ràng, build thành công và có dữ liệu seed phục vụ demo.")
    builder.heading("Hạn chế", 2)
    for item in [
        "Thanh toán QR mới được xác định là hướng phát triển, chưa triển khai callback và đối soát tự động.",
        "Document Hub/chia sẻ tài liệu học thuật chưa nằm trong phạm vi hoàn thiện của phiên bản hiện tại.",
        "Chưa có bộ automated test đầy đủ cho controller/service, hiện mới kiểm thử theo luồng chính và build.",
        "Phân tách dữ liệu đa cơ sở mới thể hiện qua routing/cơ sở mẫu, có thể mở rộng thêm tenant/branch trong database.",
    ]:
        builder.bullet(item)

    builder.page_break()
    builder.heading("CHƯƠNG 7. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)
    builder.heading("Kết quả đạt được", 2)
    builder.p("Chuyên đề đã xây dựng được hệ thống WebPhotocopyHub theo hướng ứng dụng thực tế cho dịch vụ photocopy/in ấn. Sản phẩm có public site, khu khách hàng, khu vận hành cơ sở, khu admin, dữ liệu seed, Swagger, health checks và các cơ chế kiểm soát bảo mật/cơ sở dữ liệu cần thiết. Báo cáo đã trình bày từ bối cảnh, công nghệ, yêu cầu, thiết kế đến triển khai và kiểm thử.")
    builder.heading("Giá trị thực tiễn", 2)
    builder.p("WebPhotocopyHub giúp giảm trao đổi thủ công, tăng khả năng theo dõi đơn, hỗ trợ nhân viên tập trung xử lý hàng chờ và giúp admin kiểm soát tài chính, người dùng, audit log. Đây là nền tảng có thể phát triển thành sản phẩm quản lý tiệm photocopy theo mô hình nhiều cơ sở.")
    builder.heading("Hướng phát triển", 2)
    for item in [
        "Tích hợp QR payment: sinh mã thanh toán, nhận callback ngân hàng/ví điện tử và tự động đối soát.",
        "Phát triển Document Hub: lưu trữ, phân loại và chia sẻ tài liệu học thuật có kiểm duyệt.",
        "Bổ sung notification real-time bằng SignalR cho trạng thái đơn và yêu cầu nạp tiền.",
        "Bổ sung automated tests cho service/controller và CI/CD để kiểm soát regression.",
        "Mở rộng báo cáo doanh thu, tồn kho, hiệu suất xử lý và dashboard theo từng cơ sở.",
        "Hoàn thiện triển khai production: domain, SSL, backup database, log tập trung và monitoring nâng cao.",
    ]:
        builder.bullet(item)

    builder.page_break()
    add_references(builder)

    builder.page_break()
    builder.front_heading("PHỤ LỤC")
    builder.front_subheading("Phụ lục A. Tài khoản kiểm thử")
    builder.add_table(
        "A.1",
        "Tài khoản kiểm thử dùng trong môi trường local/dev",
        ["Vai trò", "Email", "Mật khẩu", "Ghi chú"],
        [
            ["Admin", "admin@photocopyhub.local", "Admin@123456", "Quản trị toàn hệ thống"],
            ["Vận hành", "operator@photocopyhub.local", "Operator@123456", "Nhân viên vận hành cơ sở"],
            ["Khách hàng", "sinhvien01@webphotocopyhub.local hoặc sinhvien01@photocopyhub.local", "Student@123", "Khách hàng seed có dữ liệu mẫu"],
        ],
        widths=[1.05, 2.45, 1.65, 0.95],
    )
    builder.front_subheading("Phụ lục B. Endpoint kiểm tra nhanh")
    for item in [
        "Trang chủ: http://localhost:5250/Home",
        "Cơ sở Toàn Photocopy: http://localhost:5250/ToanPhotocopy",
        "Khách hàng: http://localhost:5250/ToanPhotocopy/Login",
        "ShopOperator: http://localhost:5250/ToanPhotocopy/Admin/Login",
        "Admin: http://localhost:5250/Admin/Login",
        "Swagger: http://localhost:5250/swagger",
        "Health DB: http://localhost:5250/healthz/db",
    ]:
        builder.bullet(item)

    # Formatting pass for body paragraphs.
    for p in doc.paragraphs:
        if p.style.name == "Normal" and p.text.strip():
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
        if p.style.name.startswith("Heading"):
            p.paragraph_format.keep_with_next = True

    doc.core_properties.title = TITLE
    doc.core_properties.author = STUDENT
    doc.core_properties.subject = "Chuyên đề tốt nghiệp WebPhotocopyHub"
    ensure_all_image_alt_text(doc)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_document())
